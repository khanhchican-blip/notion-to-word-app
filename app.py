import streamlit as st
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from notion_client import Client
import requests
import io
import re

# ==========================================
# 1. CẤU HÌNH API KEY CỦA BẠN Ở ĐÂY
# ==========================================
NOTION_API_KEY = st.secrets["NOTION_API_KEY"] # Thay bằng API Key của bạn
notion = Client(auth=NOTION_API_KEY)

# ==========================================
# HÀM HỖ TRỢ XỬ LÝ WORD
# ==========================================
def set_document_font(doc, font_name='Montserrat', font_size=14):
    """Cài đặt font chữ mặc định cho toàn bộ tài liệu"""
    style = doc.styles['Normal']
    font = style.font
    font.name = font_name
    font.size = Pt(font_size)

def extract_page_id(url):
    """Trích xuất Page ID từ link Notion"""
    match = re.search(r"([a-f0-9]{32})", url.replace("-", ""))
    return match.group(1) if match else None

def parse_notion_blocks(blocks, doc):
    """Hàm duyệt qua các block của Notion và đẩy vào Word"""
    for block in blocks:
        block_type = block.get('type')
        
        # 1. Xử lý Văn bản thường (Paragraph)
        if block_type == 'paragraph':
            text = "".join([t['plain_text'] for t in block['paragraph']['rich_text']])
            if text.strip():
                doc.add_paragraph(text)
                
        # 2. Xử lý Tiêu đề (Heading 1, 2, 3)
        elif block_type in ['heading_1', 'heading_2', 'heading_3']:
            level = int(block_type[-1])
            text = "".join([t['plain_text'] for t in block[block_type]['rich_text']])
            doc.add_heading(text, level=level)
            
        # 3. Xử lý Hình ảnh
        elif block_type == 'image':
            image_type = block['image']['type']
            image_url = block['image'][image_type]['url']
            try:
                # Tải ảnh từ URL của Notion
                response = requests.get(image_url)
                image_stream = io.BytesIO(response.content)
                # Chèn ảnh vào Word, chỉnh chiều rộng vừa phải để không bị tràn
                doc.add_picture(image_stream, width=Inches(6.0))
            except Exception as e:
                doc.add_paragraph(f"[Lỗi không thể tải hình ảnh: {str(e)}]")
                
        # 4. Xử lý Cột (Giả lập bằng Table ẩn viền để chống lỗi ngắt trang)
        elif block_type == 'column_list':
            # Lấy các cột con bên trong block column_list
            columns = notion.blocks.children.list(block_id=block['id']).get('results', [])
            num_cols = len(columns)
            
            if num_cols > 0:
                table = doc.add_table(rows=1, cols=num_cols)
                table.autofit = True
                
                # Duyệt qua từng cột
                for i, col in enumerate(columns):
                    col_blocks = notion.blocks.children.list(block_id=col['id']).get('results', [])
                    cell = table.cell(0, i)
                    
                    # Lấy text của từng block trong cột và gộp lại
                    cell_text = ""
                    for cb in col_blocks:
                        if cb['type'] == 'paragraph':
                            cell_text += "".join([t['plain_text'] for t in cb['paragraph']['rich_text']]) + "\n"
                        # Có thể mở rộng để xử lý ảnh/list bên trong cột ở đây
                    
                    cell.text = cell_text.strip()
            doc.add_paragraph() # Thêm dòng trống sau bảng

# ==========================================
# 2. GIAO DIỆN STREAMLIT
# ==========================================
st.set_page_config(page_title="Notion to Word", layout="centered")

# Giao diện cực kỳ đơn giản theo yêu cầu
st.title("Chuyển đổi Notion sang Word")
notion_url = st.text_input("Nhập link Notion của bạn vào đây:")

if notion_url:
    page_id = extract_page_id(notion_url)
    
    if not page_id:
        st.error("Link Notion không hợp lệ. Vui lòng kiểm tra lại.")
    else:
        if st.button("Bắt đầu xử lý"):
            with st.spinner("Đang tải dữ liệu và tạo file Word..."):
                try:
                    # Khởi tạo Word
                    doc = Document()
                    set_document_font(doc, 'Montserrat', 14)
                    
                    # Lấy tiêu đề trang (tuỳ chọn, để làm tên file)
                    page_data = notion.pages.retrieve(page_id=page_id)
                    
                    # Lấy nội dung các blocks của trang
                    blocks = notion.blocks.children.list(block_id=page_id).get('results', [])
                    
                    # Xử lý nội dung
                    parse_notion_blocks(blocks, doc)
                    
                    # Lưu file vào bộ nhớ đệm (BytesIO) để người dùng tải về
                    docx_stream = io.BytesIO()
                    doc.save(docx_stream)
                    docx_stream.seek(0)
                    
                    st.success("Tạo file thành công!")
                    
                    # Nút tải file
                    st.download_button(
                        label="Tải file Word (.docx) xuống",
                        data=docx_stream,
                        file_name="Notion_Export.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                except Exception as e:
                    st.error(f"Đã xảy ra lỗi: {str(e)}\n\n(Bạn đã Share trang này với Integration API chưa?)")